"""Build a Word document of 2026 hot pepper varieties with images.

For each pepper, fetches the lead image from Wikipedia (en.wikipedia.org REST
API). If a fetch fails (offline / blocked / no image), inserts a colored
placeholder box so the document is still complete.

Run:
    pip install python-docx requests Pillow
    python scripts/build_peppers_doc.py

Output: docs/peppers-2026.docx
"""
from __future__ import annotations

import io
import os
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

import requests
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


REPO = Path(__file__).resolve().parent.parent
OUT = REPO / "docs" / "peppers-2026.docx"
USER_AGENT = "ben-automation-system/1.0 (peppers slide builder)"


@dataclass
class Pepper:
    num: int
    he: str
    en: str
    wiki_slug: str  # English Wikipedia article slug
    image_query: str  # Commons search query (used as fallback)
    color: str  # red / yellow / orange / chocolate / purple / white
    species: str
    shu: str


PEPPERS: list[Pepper] = [
    Pepper(1,  "אנברו צהוב",            "Yellow Habanero",            "Habanero",                     "Yellow Habanero pepper",        "yellow",    "Capsicum chinense",              "100K–350K SHU"),
    Pepper(2,  "אנברו אדום",            "Red Habanero",               "Habanero",                     "Red Habanero pepper",           "red",       "Capsicum chinense",              "100K–350K SHU"),
    Pepper(3,  "קרולינה ריפר אדום",     "Carolina Reaper (Red)",      "Carolina_Reaper",              "Carolina Reaper pepper",        "red",       "Capsicum chinense",              "~1.6M–2.2M SHU"),
    Pepper(4,  "אנברו שוקולד",          "Chocolate Habanero",         "Habanero",                     "Chocolate Habanero pepper",     "chocolate", "Capsicum chinense",              "300K–425K SHU"),
    Pepper(5,  "פטליי (כתום אפריקאי)",  "Fatalii",                    "Fatalii",                      "Fatalii pepper",                "yellow",    "Capsicum chinense",              "125K–400K SHU"),
    Pepper(6,  "אנברו סגול מקומט",      "Purple Habanero",            "Habanero",                     "Purple Habanero pepper",        "purple",    "Capsicum chinense",              "100K–350K SHU"),
    Pepper(7,  "קרולינה ריפר צהוב",     "Yellow Carolina Reaper",     "Carolina_Reaper",              "Yellow Carolina Reaper pepper", "yellow",    "Capsicum chinense",              "~1.4M–2M SHU"),
    Pepper(8,  "קרולינה ריפר שוקולד",   "Chocolate Carolina Reaper",  "Carolina_Reaper",              "Chocolate Carolina Reaper",     "chocolate", "Capsicum chinense",              "~1.5M–2M SHU"),
    Pepper(9,  "גוליקה בהוט צהוב",      "Yellow Bhut Jolokia",        "Bhut_jolokia",                 "Yellow Bhut Jolokia",           "yellow",    "Capsicum chinense × frutescens", "~800K–1M SHU"),
    Pepper(10, "סקורפיו ריפר אדום",     "Trinidad Scorpion (Red)",    "Trinidad_Scorpion_pepper",     "Trinidad Scorpion pepper red",  "red",       "Capsicum chinense",              "~1.2M–2M SHU"),
    Pepper(11, "צ׳רי בומב אדום",        "Cherry Bomb (F1)",           "Cherry_pepper",                "Cherry Bomb pepper",            "red",       "Capsicum annuum",                "2.5K–5K SHU"),
    Pepper(12, "גוליקה בהוט סגול",      "Purple Bhut Jolokia",        "Bhut_jolokia",                 "Purple Bhut Jolokia",           "purple",    "Capsicum chinense",              "~800K–1M SHU"),
    Pepper(13, "ציפור תאילנדי",         "Thai Bird's Eye",            "Bird%27s_eye_chili",           "Thai bird's eye chili",         "red",       "Capsicum annuum",                "50K–100K SHU"),
    Pepper(14, "גמיקאן ילו צהוב",       "Yellow Scotch Bonnet",       "Scotch_bonnet",                "Yellow Scotch Bonnet pepper",   "yellow",    "Capsicum chinense",              "100K–350K SHU"),
    Pepper(15, "צ׳ילי תאילנדי ארוך",    "Thai Long / Spur Chili",     "Thai_pepper",                  "Thai long spur chili pepper",   "red",       "Capsicum annuum",                "15K–50K SHU"),
    Pepper(16, "עקרב לבן",              "White Trinidad Scorpion",    "Trinidad_Scorpion_pepper",     "White Trinidad Scorpion",       "white",     "Capsicum chinense",              "~1M+ SHU"),
    Pepper(17, "דובדבן אדום",           "Mexican Cherry / Pimiento",  "Pimiento",                     "Pimiento cherry pepper",        "red",       "Capsicum annuum",                "100–500 SHU"),
    Pepper(18, "קריולה סלה",            "Criolla Sella (Aji Crystal)","Capsicum_baccatum",            "Criolla Sella aji",             "yellow",    "Capsicum baccatum",              "~30K SHU"),
    Pepper(19, "אנברו מנורת לילה",      "Aji Lantern / Bishop's Crown","Bishop%27s_crown",            "Aji Lantern Bishop's crown",    "orange",    "Capsicum baccatum",              "5K–30K SHU"),
]


COLOR_HEX = {
    "red":       "FF8A8A",
    "yellow":    "FFD96A",
    "orange":    "FFA566",
    "chocolate": "C69370",
    "purple":    "C98AFF",
    "white":     "E8E8F0",
}


_DDG_SESSION: Optional[requests.Session] = None


def _ddg_session() -> requests.Session:
    global _DDG_SESSION
    if _DDG_SESSION is None:
        s = requests.Session()
        s.headers.update({
            "User-Agent": "Mozilla/5.0 (X11; Linux x86_64; rv:120.0) Gecko/20100101 Firefox/120.0",
        })
        _DDG_SESSION = s
    return _DDG_SESSION


def _ddg_vqd(query: str) -> Optional[str]:
    import re
    s = _ddg_session()
    r = s.get("https://duckduckgo.com/", params={"q": query}, timeout=10)
    if not r.ok:
        return None
    m = re.search(r"vqd=[\"']?([0-9-]+)", r.text)
    return m.group(1) if m else None


def _ddg_image_search(query: str, limit: int = 5) -> list[dict]:
    s = _ddg_session()
    vqd = _ddg_vqd(query)
    if not vqd:
        return []
    r = s.get(
        "https://duckduckgo.com/i.js",
        params={"l": "us-en", "o": "json", "q": query, "vqd": vqd, "p": "1"},
        headers={"Referer": "https://duckduckgo.com/"},
        timeout=12,
    )
    if not r.ok:
        return []
    try:
        return (r.json().get("results") or [])[:limit]
    except Exception:
        return []


_WATERMARK_DOMAINS = (
    "shutterstock.com", "dreamstime.com", "alamy.com", "istockphoto.com",
    "gettyimages.com", "depositphotos.com", "123rf.com", "stock.adobe.com",
    "agefotostock.com", "bigstockphoto.com", "stockfresh.com", "canstockphoto.com",
)


def _is_watermarked(cand: dict) -> bool:
    url = (cand.get("image") or "").lower()
    src = (cand.get("source") or cand.get("url") or "").lower()
    haystack = f"{url} {src}"
    return any(d in haystack for d in _WATERMARK_DOMAINS)


def fetch_image(p: Pepper) -> Optional[bytes]:
    """Search DuckDuckGo Images for the pepper, fetch first viable result.

    Skips known watermarked stock-photo sources. Uses the Bing CDN thumbnail
    (always reachable) as the actual fetch URL.
    """
    candidates = _ddg_image_search(f"{p.image_query} pepper", limit=15)
    for cand in candidates:
        if _is_watermarked(cand):
            continue
        for url in (cand.get("thumbnail"), cand.get("image")):
            if not url:
                continue
            try:
                ir = requests.get(
                    url,
                    headers={"User-Agent": "Mozilla/5.0", "Referer": "https://duckduckgo.com/"},
                    timeout=12,
                )
                ct = ir.headers.get("content-type", "")
                if ir.ok and ct.startswith("image/") and len(ir.content) > 5000:
                    return ir.content
            except Exception:
                continue

    # Fallback: Wikipedia REST summary (works only when not egress-blocked)
    try:
        r = requests.get(
            f"https://en.wikipedia.org/api/rest_v1/page/summary/{p.wiki_slug}",
            headers={"User-Agent": USER_AGENT, "Accept": "application/json"},
            timeout=8,
        )
        if r.ok:
            j = r.json()
            src = (j.get("originalimage") or {}).get("source") or (j.get("thumbnail") or {}).get("source")
            if src:
                ir = requests.get(src, headers={"User-Agent": USER_AGENT}, timeout=12)
                if ir.ok and ir.headers.get("content-type", "").startswith("image/"):
                    return ir.content
    except Exception:
        pass

    return None


def normalize_for_docx(image_bytes: bytes, max_side: int = 1200) -> bytes:
    """Open with PIL, re-encode as JPEG with explicit DPI metadata.
    python-docx needs DPI info to compute physical size; many web JPEGs/WebPs
    omit DPI which causes a division-by-zero in add_picture.
    """
    from PIL import Image
    im = Image.open(io.BytesIO(image_bytes))
    if im.mode in ("RGBA", "P", "LA"):
        bg = Image.new("RGB", im.size, (255, 255, 255))
        bg.paste(im, mask=im.split()[-1] if im.mode in ("RGBA", "LA") else None)
        im = bg
    elif im.mode != "RGB":
        im = im.convert("RGB")
    if max(im.size) > max_side:
        im.thumbnail((max_side, max_side), Image.LANCZOS)
    out = io.BytesIO()
    im.save(out, format="JPEG", quality=88, dpi=(96, 96), optimize=True)
    return out.getvalue()


def build_placeholder_png(color_hex: str, size: int = 600) -> bytes:
    """Build a simple solid-color PNG as fallback when image fetch fails.
    Avoids requiring Pillow at runtime by hand-rolling a tiny PNG.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
        rgb = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        im = Image.new("RGB", (size, size), rgb)
        d = ImageDraw.Draw(im)
        d.rectangle([0, 0, size - 1, size - 1], outline=(60, 60, 60), width=4)
        try:
            font = ImageFont.truetype("DejaVuSans-Bold.ttf", 28)
        except Exception:
            font = ImageFont.load_default()
        msg = "image unavailable"
        bbox = d.textbbox((0, 0), msg, font=font)
        w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
        d.text(((size - w) / 2, (size - h) / 2), msg, fill=(40, 40, 40), font=font)
        buf = io.BytesIO()
        im.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        # 1x1 PNG of the color, scaled by Word
        import struct, zlib
        rgb = bytes(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
        sig = b"\x89PNG\r\n\x1a\n"
        ihdr = struct.pack(">IIBBBBB", size, size, 8, 2, 0, 0, 0)
        ihdr_chunk = b"IHDR" + ihdr
        ihdr_full = struct.pack(">I", len(ihdr)) + ihdr_chunk + struct.pack(">I", zlib.crc32(ihdr_chunk))
        raw = b""
        for _ in range(size):
            raw += b"\x00" + (rgb * size)
        idat_data = zlib.compress(raw, 9)
        idat_chunk = b"IDAT" + idat_data
        idat = struct.pack(">I", len(idat_data)) + idat_chunk + struct.pack(">I", zlib.crc32(idat_chunk))
        iend_chunk = b"IEND"
        iend = struct.pack(">I", 0) + iend_chunk + struct.pack(">I", zlib.crc32(iend_chunk))
        return sig + ihdr_full + idat + iend


def set_run_complex_script(run, size_pt: int, bold: bool = False, rtl: bool = True) -> None:
    """Apply size to both Latin (w:sz) and complex-script (w:szCs) so Hebrew
    text renders at the requested size, plus optional RTL flag and bold-cs.
    """
    rPr = run._element.get_or_add_rPr()
    half_pts = str(size_pt * 2)
    for tag in ("w:sz", "w:szCs"):
        for existing in rPr.findall(qn(tag)):
            rPr.remove(existing)
        el = OxmlElement(tag)
        el.set(qn("w:val"), half_pts)
        rPr.append(el)
    if bold:
        for tag in ("w:b", "w:bCs"):
            if not rPr.findall(qn(tag)):
                el = OxmlElement(tag)
                el.set(qn("w:val"), "1")
                rPr.append(el)
    if rtl:
        for existing in rPr.findall(qn("w:rtl")):
            rPr.remove(existing)
        el = OxmlElement("w:rtl")
        el.set(qn("w:val"), "1")
        rPr.append(el)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex: str = "1A2236", size: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tc_pr.append(borders)


def add_pepper_slide(doc: Document, p: Pepper, image_bytes: Optional[bytes], used_fallback: bool) -> None:
    """One pepper per page: title + English + meta on the right (RTL), image on the left."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Cm(12)
    table.columns[1].width = Cm(13)

    # Right column: text (RTL — first column visually = right side in RTL flow)
    right_cell = table.cell(0, 0)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_borders(right_cell, "FFFFFF", size=0)

    num_p = right_cell.paragraphs[0]
    num_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    num_run = num_p.add_run(f"{p.num} / {len(PEPPERS)}")
    num_run.font.size = Pt(12)
    num_run.font.color.rgb = RGBColor(0x90, 0x90, 0xA0)

    he_p = right_cell.add_paragraph()
    he_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    he_p.paragraph_format.space_before = Pt(8)
    he_run = he_p.add_run(p.he)
    he_run.bold = True
    he_run.font.color.rgb = RGBColor(0x10, 0x10, 0x18)
    set_run_complex_script(he_run, size_pt=36, bold=True, rtl=True)

    en_p = right_cell.add_paragraph()
    en_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    en_run = en_p.add_run(p.en)
    en_run.italic = True
    en_run.font.size = Pt(18)
    en_run.font.color.rgb = RGBColor(0x55, 0x55, 0x66)

    meta_p = right_cell.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_p.paragraph_format.space_before = Pt(16)
    for i, label in enumerate([p.species, p.shu]):
        if i:
            meta_p.add_run("   ·   ").font.color.rgb = RGBColor(0xAA, 0xAA, 0xB8)
        r = meta_p.add_run(label)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x66)

    if used_fallback:
        note_p = right_cell.add_paragraph()
        note_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        note_p.paragraph_format.space_before = Pt(20)
        note_r = note_p.add_run(
            "תמונה לא נטענה — ראה: "
            f"https://en.wikipedia.org/wiki/{p.wiki_slug}"
        )
        note_r.font.size = Pt(9)
        note_r.font.color.rgb = RGBColor(0xC9, 0x30, 0x30)

    # Left column: image
    left_cell = table.cell(0, 1)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_borders(left_cell, "1A2236", size=12)

    img_p = left_cell.paragraphs[0]
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_bytes:
        normalized = normalize_for_docx(image_bytes)
        try:
            img_p.add_run().add_picture(io.BytesIO(normalized), width=Cm(12))
        except Exception as e:
            print(f"  [{p.num}] add_picture failed: {e}, using placeholder", file=sys.stderr)
            img_p.add_run().add_picture(io.BytesIO(build_placeholder_png(COLOR_HEX[p.color])), width=Cm(12))
    else:
        img_p.add_run().add_picture(io.BytesIO(build_placeholder_png(COLOR_HEX[p.color])), width=Cm(12))

    # Page break after each slide except last
    if p.num < len(PEPPERS):
        doc.add_page_break()


def add_title_page(doc: Document) -> None:
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_r = title_p.add_run("זני פלפלים חריפים")
    title_r.bold = True
    title_r.font.color.rgb = RGBColor(0xC9, 0x30, 0x30)
    set_run_complex_script(title_r, size_pt=56, bold=True, rtl=True)

    year_p = doc.add_paragraph()
    year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year_p.paragraph_format.space_before = Pt(20)
    year_r = year_p.add_run("2026")
    year_r.bold = True
    year_r.font.size = Pt(48)
    year_r.font.color.rgb = RGBColor(0x66, 0x66, 0x80)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(40)
    sub_r = sub_p.add_run(f"{len(PEPPERS)} זנים — שמות עבריים, אנגליים ותמונות")
    sub_r.font.size = Pt(14)
    sub_r.font.color.rgb = RGBColor(0x80, 0x80, 0x90)
    set_run_complex_script(sub_r, size_pt=14, bold=False, rtl=True)

    doc.add_page_break()


def main() -> int:
    doc = Document()

    # Landscape A4
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(2)
    section.top_margin = section.bottom_margin = Cm(1.5)

    add_title_page(doc)

    fetched = 0
    for p in PEPPERS:
        print(f"[{p.num}/{len(PEPPERS)}] {p.he} ({p.en})…", file=sys.stderr)
        img = fetch_image(p)
        used_fallback = img is None
        if img:
            fetched += 1
        add_pepper_slide(doc, p, img, used_fallback)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"\nWrote {OUT} ({fetched}/{len(PEPPERS)} images fetched, "
          f"{len(PEPPERS) - fetched} placeholders)", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
