"""Build pepper sticker sheets from grandpa's hand-shot photos.

Uses the same sticker composition / page layout as build_peppers_stickers.py,
but sources photos from assets/pepper-photos/ with a hand-curated filename →
Hebrew name mapping (because OCR on user-shot photos is unreliable).

Output:
    docs/peppers-grandpa-stickers.pdf
    docs/peppers-grandpa-stickers.docx
"""
from __future__ import annotations

import io
import sys
from pathlib import Path
from typing import Optional

from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_peppers_stickers import (  # type: ignore
    CELL_W, CELL_H, COLS, ROWS, DPI, LOGO_TRANS,
    compose_page, make_sticker,
)

REPO = Path(__file__).resolve().parent.parent
PHOTOS_DIR = REPO / "assets" / "pepper-photos"
DOCS = REPO / "docs"

# Hand-curated mapping (identified by visual inspection of each file).
# Skips 2756fe97 — unlabeled photo per user instruction.
PHOTO_MAP: list[tuple[str, str]] = [
    ("07cc495a-d456-4f56-8d4e-1b0ac62fd861.jpeg", "7 פוד לוסי"),
    ("0fc2ba22-d6e8-4d9e-b9b4-53e6967ecb7f.jpeg", "הבנרו אפרסק"),
    ("1a044bea-7b16-42ce-aadb-980ba7ef9f55.jpeg", "ג'מאיקאן מאשרום צהוב"),
    ("28046bf4-8588-467c-846a-5f2161a07705.jpeg", "הבנרו צהוב"),
    ("42fbf3e4-ea24-4458-9469-30b1210f3ab5.jpeg", "אחי לימון דרופ"),
    ("450f15a9-15ac-4b3e-8422-86cd8a163345.jpeg", "סקורפיון לבן"),
    ("48444c0b-83b9-4aba-bedf-4f4049e71216.jpeg", "פפרוני פפר"),
    ("4fb120bb-9f56-4505-ab20-8114085aba09.jpeg", "סקוטש בונט"),
    ("6620ac26-4195-42b5-b061-95d36373a792.jpeg", "הבנרו אדום"),
    ("693253d2-e93c-4f83-b881-2fffddbdf6b7.jpeg", "7 פוד ריפר שוקולד"),
    ("69a42848-1e6b-4ac4-93f4-403ec6d95d0d.jpeg", "גוסט פפר שוקולד"),
    ("6bf108e4-be02-4762-8ef5-994ac0a1efe0.jpeg", "מורוגה אדום"),
    ("7aa06236-8d85-42ca-8652-fd85d7860fe8.jpeg", "דארסט נאגא"),
    ("7b8e0dd2-f6ab-45a1-82f0-72f35d669928.jpeg", "חלפיניו"),
    ("9beb8127-e55a-45a0-930b-da296337f4b7.jpeg", "נאגא מוריץ'"),
    ("9cd43c69-0178-41db-b7c9-144542423776.jpeg", "קוסטה ריקה צ'ילי"),
    ("a0b98327-eea0-47d1-b12c-9ee371520107.jpeg", "הוט צ'ילי פפר"),
    ("a2487a88-015b-4bc6-9fef-6c32e55e4d4e.jpeg", "פתאלי גורמה אדום"),
    ("a679289e-db2b-4c57-b0b4-80f23e871d56.jpeg", "אפוקאליפסה סקורפיון שוקולד"),
    ("baefb719-0245-42b0-93a2-028a4080ed69.jpeg", "הבנרו מנורת נייר"),
    ("bb3d5354-6a95-47b5-b78f-576cb84cefe5.jpeg", "קאיין ארוך"),
    ("d166dc5f-5cd0-4761-819c-d527903652d4.jpeg", "גאיי'ס גוסט סקורפיון אפרסק"),
    ("df813141-8c22-4736-9121-d34477b0c866.jpeg", "קרולינה ריפר"),
    ("e73896f6-b29a-4648-8741-2a828d52158d.jpeg", "גוסט פפר צהוב"),
    ("f6eec934-ac06-47a1-9649-b1c0f0e145a0.jpeg", "הבנרו לבן"),
    ("fe3a08ed-9daa-4eea-bcfb-0b3f3438c4d8.jpeg", "סופר גוסט פפר"),
]

# Files that are screenshots and need top/bottom UI cropped before use.
# Values: (top_fraction_to_remove, bottom_fraction_to_remove)
CROP_OVERRIDES: dict[str, tuple[float, float]] = {
    "a2487a88-015b-4bc6-9fef-6c32e55e4d4e.jpeg": (0.21, 0.10),
}


def load_photo(fname: str) -> Optional[bytes]:
    path = PHOTOS_DIR / fname
    if not path.exists():
        print(f"  MISSING: {fname}", file=sys.stderr)
        return None
    im = Image.open(path).convert("RGB")
    # Crop screenshot UI if needed
    if fname in CROP_OVERRIDES:
        top_f, bot_f = CROP_OVERRIDES[fname]
        w, h = im.size
        im = im.crop((0, int(h * top_f), w, int(h * (1 - bot_f))))
    buf = io.BytesIO()
    im.save(buf, format="JPEG", quality=92)
    return buf.getvalue()


def main() -> int:
    if not LOGO_TRANS.exists():
        print(f"ERROR: missing {LOGO_TRANS}", file=sys.stderr)
        return 1
    logo = Image.open(LOGO_TRANS).convert("RGBA")
    print(f"Page: A4 portrait @ {DPI}dpi, {COLS}x{ROWS} grid", file=sys.stderr)
    print(f"Cell: {CELL_W}x{CELL_H}px", file=sys.stderr)
    print(f"Stickers to build: {len(PHOTO_MAP)}", file=sys.stderr)

    stickers: list[Image.Image] = []
    for fname, name_he in PHOTO_MAP:
        img = load_photo(fname)
        if not img:
            continue
        print(f"  -> {name_he}", file=sys.stderr)
        s = make_sticker(img, name_he, logo, (CELL_W, CELL_H))
        stickers.append(s)

    if not stickers:
        print("No stickers built", file=sys.stderr)
        return 2

    per_page = COLS * ROWS
    pages: list[Image.Image] = []
    for i in range(0, len(stickers), per_page):
        pages.append(compose_page(stickers[i:i + per_page]))

    DOCS.mkdir(parents=True, exist_ok=True)
    pdf_path = DOCS / "peppers-grandpa-stickers.pdf"
    pages[0].save(
        pdf_path, "PDF", resolution=DPI,
        save_all=True, append_images=pages[1:] if len(pages) > 1 else [],
    )
    print(f"\nWrote {pdf_path}  ({len(pages)} pages, {len(stickers)} stickers)",
          file=sys.stderr)

    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.shared import Cm
        doc = Document()
        s = doc.sections[0]
        s.orientation = WD_ORIENT.PORTRAIT
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.left_margin = s.right_margin = Cm(0)
        s.top_margin = s.bottom_margin = Cm(0)
        for idx, page in enumerate(pages):
            buf = io.BytesIO()
            page.save(buf, format="PNG", dpi=(DPI, DPI))
            buf.seek(0)
            if idx == 0 and doc.paragraphs:
                p = doc.paragraphs[0]
            else:
                if idx:
                    doc.add_page_break()
                p = doc.add_paragraph()
            p.paragraph_format.space_before = Cm(0)
            p.paragraph_format.space_after = Cm(0)
            p.add_run().add_picture(buf, width=Cm(21.0))
        docx_path = DOCS / "peppers-grandpa-stickers.docx"
        doc.save(docx_path)
        print(f"Wrote {docx_path}", file=sys.stderr)
    except Exception as e:
        print(f"DOCX export skipped: {e}", file=sys.stderr)

    return 0


if __name__ == "__main__":
    sys.exit(main())
