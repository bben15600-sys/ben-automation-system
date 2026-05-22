"""Build a printable PDF of pepper stickers — 12 stickers per A4 page.

Each sticker is an oval with the pepper image as background, the Hebrew name
overlaid at the bottom, and the brand logo (no background) in the top corner.

Reuses the canonical PEPPERS list and fetch_image() from build_peppers_doc.

Run:
    pip install python-docx requests Pillow rembg onnxruntime
    python scripts/build_peppers_stickers.py

Outputs:
    docs/peppers-2026-stickers.pdf
    docs/peppers-2026-stickers.docx
"""
from __future__ import annotations

import io
import os
import sys
from pathlib import Path
from typing import Optional

from PIL import Image, ImageDraw, ImageFilter, ImageFont

sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_peppers_doc import PEPPERS, fetch_image as _fetch_image  # type: ignore

REPO = Path(__file__).resolve().parent.parent
ASSETS = REPO / "assets"
CACHE = ASSETS / ".image-cache"
DOCS = REPO / "docs"
LOGO_TRANS = ASSETS / "logo-nobg.png"


def fetch_image(pepper) -> Optional[bytes]:
    """Cached version of build_peppers_doc.fetch_image, resilient to
    network errors."""
    CACHE.mkdir(parents=True, exist_ok=True)
    cached = CACHE / f"pepper-{pepper.num:02d}.jpg"
    if cached.exists() and cached.stat().st_size > 1000:
        return cached.read_bytes()
    for attempt in range(3):
        try:
            data = _fetch_image(pepper)
            if data:
                cached.write_bytes(data)
                return data
            return None
        except Exception as e:
            print(f"  fetch retry {attempt + 1}/3 for {pepper.he}: {e}", file=sys.stderr)
    return None

# A4 @ 300 DPI portrait
DPI = 300
PAGE_W_MM, PAGE_H_MM = 210, 297

def mm(x: float) -> int:
    return round(x / 25.4 * DPI)

PAGE_W, PAGE_H = mm(PAGE_W_MM), mm(PAGE_H_MM)

# Layout: 3 cols × 4 rows = 12 per page
COLS, ROWS = 3, 4
M_H, M_V = mm(11), mm(14)     # page margins
GAP_H, GAP_V = mm(5), mm(5)   # between stickers

CELL_W = (PAGE_W - 2 * M_H - (COLS - 1) * GAP_H) // COLS
CELL_H = (PAGE_H - 2 * M_V - (ROWS - 1) * GAP_V) // ROWS

# Font paths
FONT_HEB_BOLD = "/usr/share/fonts/truetype/noto/NotoSansHebrew-Bold.ttf"
FONT_HEB = "/usr/share/fonts/truetype/noto/NotoSansHebrew-Regular.ttf"
FONT_LATIN = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"


def load_font(path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(path, size)


def measure_text_width(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont) -> int:
    """Visual width of text. Uses bbox(direction=rtl) to capture true extent
    (including marks/diacritics) rather than just the advance-width."""
    try:
        bbox = draw.textbbox((0, 0), text, font=font, direction="rtl")
    except TypeError:
        bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0]


def fit_text_to_width(draw: ImageDraw.ImageDraw, text: str, font_path: str,
                      max_w: int, start_size: int, min_size: int = 14) -> ImageFont.FreeTypeFont:
    size = start_size
    while size >= min_size:
        f = load_font(font_path, size)
        if measure_text_width(draw, text, f) <= max_w:
            return f
        size -= 2
    return load_font(font_path, min_size)


def crop_cover(im: Image.Image, target_w: int, target_h: int) -> Image.Image:
    """Resize+crop image to exactly target dims, preserving aspect (cover)."""
    src_w, src_h = im.size
    scale = max(target_w / src_w, target_h / src_h)
    new_w, new_h = int(src_w * scale), int(src_h * scale)
    im = im.resize((new_w, new_h), Image.LANCZOS)
    left = (new_w - target_w) // 2
    top = (new_h - target_h) // 2
    return im.crop((left, top, left + target_w, top + target_h))


def make_sticker(pepper_img_bytes: bytes, name_he: str, logo: Image.Image,
                 size: tuple[int, int]) -> Image.Image:
    """Compose one rounded-square sticker."""
    sw, sh = size
    canvas = Image.new("RGBA", (sw, sh), (0, 0, 0, 0))

    # Rounded-rectangle mask (sticker shape)
    radius = max(12, int(min(sw, sh) * 0.08))
    mask = Image.new("L", (sw, sh), 0)
    ImageDraw.Draw(mask).rounded_rectangle(
        (0, 0, sw - 1, sh - 1), radius=radius, fill=255
    )

    # Pepper image cropped to sticker shape
    pep = Image.open(io.BytesIO(pepper_img_bytes)).convert("RGB")
    pep = crop_cover(pep, sw, sh)
    canvas.paste(pep, (0, 0), mask=mask)

    # Bottom name banner (semi-transparent gradient)
    band_h = int(sh * 0.26)
    band = Image.new("RGBA", (sw, band_h), (0, 0, 0, 0))
    bdraw = ImageDraw.Draw(band)
    for y in range(band_h):
        alpha = int(220 * (y / band_h) ** 0.8 + 35)
        bdraw.rectangle((0, y, sw, y + 1), fill=(0, 0, 0, min(alpha, 220)))
    band_mask = mask.crop((0, sh - band_h, sw, sh))
    canvas.paste(band, (0, sh - band_h), mask=band_mask)

    d = ImageDraw.Draw(canvas)

    # Hebrew name
    max_text_w = int(sw * 0.82)
    name_font = fit_text_to_width(d, name_he, FONT_HEB_BOLD, max_text_w,
                                  start_size=int(sh * 0.11), min_size=16)
    try:
        bbox = d.textbbox((0, 0), name_he, font=name_font, direction="rtl")
    except TypeError:
        bbox = d.textbbox((0, 0), name_he, font=name_font)
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    tx = (sw - tw) // 2 - bbox[0]
    ty = sh - band_h + (band_h - th) // 2 - bbox[1] - int(sh * 0.02)
    # Text shadow for legibility (multi-direction)
    for ox, oy in ((-2, 0), (2, 0), (0, -2), (0, 2), (-1, -1), (1, 1), (-1, 1), (1, -1)):
        try:
            d.text((tx + ox, ty + oy), name_he, font=name_font,
                   fill=(0, 0, 0, 220), direction="rtl")
        except TypeError:
            d.text((tx + ox, ty + oy), name_he, font=name_font, fill=(0, 0, 0, 220))
    try:
        d.text((tx, ty), name_he, font=name_font,
               fill=(255, 255, 255, 255), direction="rtl")
    except TypeError:
        d.text((tx, ty), name_he, font=name_font, fill=(255, 255, 255, 255))

    # Logo top-right corner
    logo_size = int(sw * 0.28)
    logo_resized = logo.copy()
    logo_resized.thumbnail((logo_size, logo_size), Image.LANCZOS)
    pad = int(sw * 0.045)
    lx = sw - logo_resized.width - pad
    ly = pad
    # Soft white glow behind logo so it pops on busy backgrounds
    glow = Image.new("RGBA", logo_resized.size, (0, 0, 0, 0))
    gd = ImageDraw.Draw(glow)
    gd.ellipse((0, 0, glow.width - 1, glow.height - 1), fill=(255, 255, 255, 150))
    glow = glow.filter(ImageFilter.GaussianBlur(radius=8))
    canvas.alpha_composite(glow, (lx, ly))
    canvas.alpha_composite(logo_resized, (lx, ly))

    # Thin red outline along the rounded-square edge
    outline = Image.new("RGBA", (sw, sh), (0, 0, 0, 0))
    od = ImageDraw.Draw(outline)
    od.rounded_rectangle(
        (1, 1, sw - 2, sh - 2),
        radius=radius,
        outline=(201, 48, 48, 255),
        width=max(3, int(sw * 0.012)),
    )
    canvas.alpha_composite(outline)

    return canvas


def compose_page(stickers: list[Image.Image]) -> Image.Image:
    page = Image.new("RGB", (PAGE_W, PAGE_H), (255, 255, 255))
    idx = 0
    for r in range(ROWS):
        for c in range(COLS):
            if idx >= len(stickers):
                return page
            x = M_H + c * (CELL_W + GAP_H)
            y = M_V + r * (CELL_H + GAP_V)
            s = stickers[idx]
            if s.size != (CELL_W, CELL_H):
                s = s.resize((CELL_W, CELL_H), Image.LANCZOS)
            page.paste(s, (x, y), mask=s.split()[-1])
            idx += 1
    return page


def main() -> int:
    if not LOGO_TRANS.exists():
        print(f"ERROR: missing {LOGO_TRANS}. Run logo background removal first.",
              file=sys.stderr)
        return 1
    logo = Image.open(LOGO_TRANS).convert("RGBA")
    print(f"Loaded logo: {logo.size}", file=sys.stderr)
    print(f"Page: {PAGE_W}x{PAGE_H}px ({PAGE_W_MM}x{PAGE_H_MM}mm @ {DPI}dpi)",
          file=sys.stderr)
    print(f"Cell: {CELL_W}x{CELL_H}px (~{CELL_W*25.4/DPI:.0f}x{CELL_H*25.4/DPI:.0f}mm)",
          file=sys.stderr)

    stickers: list[Image.Image] = []
    for p in PEPPERS:
        print(f"[{p.num}/{len(PEPPERS)}] {p.he}…", file=sys.stderr)
        img_bytes = fetch_image(p)
        if not img_bytes:
            print(f"  no image for {p.he}, skipping", file=sys.stderr)
            continue
        s = make_sticker(img_bytes, p.he, logo, (CELL_W, CELL_H))
        stickers.append(s)

    if not stickers:
        print("No stickers produced", file=sys.stderr)
        return 2

    # Compose pages of 12
    per_page = COLS * ROWS
    pages: list[Image.Image] = []
    for i in range(0, len(stickers), per_page):
        pages.append(compose_page(stickers[i:i + per_page]))

    # PDF
    DOCS.mkdir(parents=True, exist_ok=True)
    pdf_path = DOCS / "peppers-2026-stickers.pdf"
    pages[0].save(
        pdf_path, "PDF", resolution=DPI,
        save_all=True, append_images=pages[1:] if len(pages) > 1 else [],
    )
    print(f"\nWrote {pdf_path}  ({len(pages)} pages, {len(stickers)} stickers)",
          file=sys.stderr)

    # DOCX (each page as a full-bleed image on its own page)
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.shared import Cm
        doc = Document()
        s = doc.sections[0]
        s.orientation = WD_ORIENT.PORTRAIT
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.left_margin = s.right_margin = Cm(0)
        s.top_margin = s.bottom_margin = Cm(0)
        for idx, page in enumerate(pages):
            buf = io.BytesIO()
            page.save(buf, format="PNG", dpi=(DPI, DPI))
            buf.seek(0)
            if idx == 0 and doc.paragraphs:
                p = doc.paragraphs[0]
            else:
                if idx:
                    doc.add_page_break()
                p = doc.add_paragraph()
            p.paragraph_format.space_before = Cm(0)
            p.paragraph_format.space_after = Cm(0)
            p.add_run().add_picture(buf, width=Cm(21.0))
        docx_path = DOCS / "peppers-2026-stickers.docx"
        doc.save(docx_path)
        print(f"Wrote {docx_path}", file=sys.stderr)
    except Exception as e:
        print(f"DOCX export skipped: {e}", file=sys.stderr)

    return 0


if __name__ == "__main__":
    sys.exit(main())
