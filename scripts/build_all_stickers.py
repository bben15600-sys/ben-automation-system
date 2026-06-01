"""Combined sticker sheet: 20 curated web-image peppers + 26 grandpa photos.

Uses the same sticker design as build_peppers_stickers / build_grandpa_stickers
(rounded-square, 12 per A4 portrait, Hebrew name banner, brand logo top-right).

Output:
    docs/peppers-all-stickers.pdf
    docs/peppers-all-stickers.docx
"""
from __future__ import annotations

import io
import sys
from pathlib import Path
from typing import Optional

from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_peppers_stickers import (  # type: ignore
    CELL_W, CELL_H, COLS, ROWS, DPI, LOGO_TRANS,
    compose_page, make_sticker,
    fetch_image as fetch_web_image,
)
from build_peppers_doc import PEPPERS  # type: ignore
from build_grandpa_stickers import PHOTO_MAP, PHOTOS_DIR, CROP_OVERRIDES  # type: ignore

REPO = Path(__file__).resolve().parent.parent
DOCS = REPO / "docs"


def load_local_photo(fname: str) -> Optional[bytes]:
    path = PHOTOS_DIR / fname
    if not path.exists():
        return None
    im = Image.open(path).convert("RGB")
    if fname in CROP_OVERRIDES:
        top_f, bot_f = CROP_OVERRIDES[fname]
        w, h = im.size
        im = im.crop((0, int(h * top_f), w, int(h * (1 - bot_f))))
    buf = io.BytesIO()
    im.save(buf, format="JPEG", quality=92)
    return buf.getvalue()


def main() -> int:
    if not LOGO_TRANS.exists():
        print(f"ERROR: missing {LOGO_TRANS}", file=sys.stderr)
        return 1
    logo = Image.open(LOGO_TRANS).convert("RGBA")
    total = len(PEPPERS) + len(PHOTO_MAP)
    print(f"Building {total} stickers: {len(PEPPERS)} web + {len(PHOTO_MAP)} local",
          file=sys.stderr)

    stickers: list[Image.Image] = []

    # Phase 1: web-image peppers (uses image cache from build_peppers_stickers)
    for p in PEPPERS:
        print(f"  [web {p.num}] {p.he}…", file=sys.stderr)
        img = fetch_web_image(p)
        if not img:
            print(f"    no image, skipping", file=sys.stderr)
            continue
        s = make_sticker(img, p.he, logo, (CELL_W, CELL_H))
        stickers.append(s)

    # Phase 2: grandpa's local photos (uses pepper-photos-clean folder)
    for idx, (fname, name_he) in enumerate(PHOTO_MAP, 1):
        print(f"  [local {idx}] {name_he}…", file=sys.stderr)
        img = load_local_photo(fname)
        if not img:
            print(f"    missing {fname}, skipping", file=sys.stderr)
            continue
        s = make_sticker(img, name_he, logo, (CELL_W, CELL_H))
        stickers.append(s)

    if not stickers:
        print("No stickers built", file=sys.stderr)
        return 2

    per_page = COLS * ROWS
    pages: list[Image.Image] = []
    for i in range(0, len(stickers), per_page):
        pages.append(compose_page(stickers[i:i + per_page]))

    DOCS.mkdir(parents=True, exist_ok=True)
    pdf_path = DOCS / "peppers-all-stickers.pdf"
    pages[0].save(
        pdf_path, "PDF", resolution=DPI,
        save_all=True, append_images=pages[1:] if len(pages) > 1 else [],
    )
    print(f"\nWrote {pdf_path}  ({len(pages)} pages, {len(stickers)} stickers)",
          file=sys.stderr)

    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.shared import Cm
        doc = Document()
        s = doc.sections[0]
        s.orientation = WD_ORIENT.PORTRAIT
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.left_margin = s.right_margin = Cm(0)
        s.top_margin = s.bottom_margin = Cm(0)
        for idx, page in enumerate(pages):
            buf = io.BytesIO()
            page.save(buf, format="PNG", dpi=(DPI, DPI))
            buf.seek(0)
            if idx == 0 and doc.paragraphs:
                p = doc.paragraphs[0]
            else:
                if idx:
                    doc.add_page_break()
                p = doc.add_paragraph()
            p.paragraph_format.space_before = Cm(0)
            p.paragraph_format.space_after = Cm(0)
            p.add_run().add_picture(buf, width=Cm(21.0))
        docx_path = DOCS / "peppers-all-stickers.docx"
        doc.save(docx_path)
        print(f"Wrote {docx_path}", file=sys.stderr)
    except Exception as e:
        print(f"DOCX export skipped: {e}", file=sys.stderr)

    return 0


if __name__ == "__main__":
    sys.exit(main())
